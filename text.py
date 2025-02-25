import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
import time
from docx import Document

# Load environment variables
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Load Gemini AI Model
model = genai.GenerativeModel('gemini-1.5-flash')

# Indian Hospital Emergency Contact Numbers
INDIAN_HOSPITALS = {
    "AIIMS Delhi": "+91-11-26588500",
    "Apollo Hospital, Chennai": "+91-44-28293333",
    "Fortis Hospital, Mumbai": "+91-22-43654365",
    "Medanta Hospital, Gurugram": "+91-124-4141414",
    "CMC Vellore": "+91-416-2281000",
}

# Function to get AI diagnosis
def get_diagnosis(name, age, gender, symptoms, history):
    input_prompt = f"""
    You are an AI medical assistant. A patient has provided the following details:

    👤 **Patient Name:** {name}
    🎂 **Age:** {age}
    🚻 **Gender:** {gender}
    🔍 **Symptoms:** {symptoms}
    📜 **Past Medical History:** {history if history else "None"}

    Your task:
    1️⃣ **Diagnose** the possible disease(s) based on symptoms.
    2️⃣ **Explain Causes** of the disease.
    3️⃣ **Suggest Treatments** (medications, home remedies, and medical procedures).
    4️⃣ **Advise Next Steps**, such as consulting a doctor or lifestyle changes.
    5️⃣ **If the condition is critical, provide emergency alert and list top Indian hospitals.**

    ⚠️ **Important**: If symptoms indicate a life-threatening condition, warn the user and suggest immediate medical assistance.
    """

    response = model.generate_content(input_prompt)
    return response.text

# Function to generate a medical report (Word file)
def generate_report(name, age, gender, symptoms, history, diagnosis):
    doc = Document()
    doc.add_heading('🩺 Medical Diagnosis Report', level=1)

    doc.add_paragraph(f"👤 **Patient Name:** {name}")
    doc.add_paragraph(f"🎂 **Age:** {age}")
    doc.add_paragraph(f"🚻 **Gender:** {gender}")
    doc.add_paragraph(f"🔍 **Symptoms:** {symptoms}")
    doc.add_paragraph(f"📜 **Past Medical History:** {history if history else 'None'}")

    doc.add_paragraph("\n📢 **AI Diagnosis & Treatment**")
    doc.add_paragraph(diagnosis)

    # Check for critical cases
    CRITICAL_KEYWORDS = ["stroke", "heart attack", "severe breathing difficulty", "unconscious", "chest pain"]
    
    if any(keyword in diagnosis.lower() for keyword in CRITICAL_KEYWORDS):
        doc.add_paragraph("\n🚨 **CRITICAL CONDITION DETECTED! SEEK IMMEDIATE MEDICAL HELP!**", style='Intense Quote')
        doc.add_paragraph("\n🏥 **Emergency Contact Numbers (India)**")
        for hospital, contact in INDIAN_HOSPITALS.items():
            doc.add_paragraph(f"📌 {hospital}: {contact}")

    # Save the document
    file_path = f"{name}_Medical_Report.docx"
    doc.save(file_path)
    return file_path

# Streamlit App Configuration
st.set_page_config(page_title="🩺 AI Medical Diagnosis Chatbot", layout="wide")

st.title("🤖 AI-Powered Medical Diagnosis")

# Initialize session state for chat history and timeout
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'last_interaction_time' not in st.session_state:
    st.session_state.last_interaction_time = time.time()

# Session timeout (2 minutes)
if time.time() - st.session_state.last_interaction_time > 120:
    st.session_state.chat_history = []

# Creating input form
with st.form("medical_form"):
    col1, col2 = st.columns(2)
    
    with col1:
        name = st.text_input("👤 Patient Name", placeholder="Enter your full name")
        age = st.number_input("🎂 Age", min_value=1, max_value=120, value=25)
        gender = st.selectbox("🚻 Gender", ["Male", "Female", "Other"])

    with col2:
        symptoms = st.text_area("🔍 Enter Your Symptoms", placeholder="E.g., fever, cough, headache, body pain")
        past_medical_history = st.text_area("📜 Do you have any past medical conditions? (Optional)", 
                                            placeholder="E.g., Diabetes, Hypertension, Asthma")

    submit = st.form_submit_button("🔍 Get Diagnosis")

# Process request on button click
if submit and symptoms.strip() and name.strip():
    st.session_state.last_interaction_time = time.time()  # Update session timer

    # Generate AI diagnosis
    response = get_diagnosis(name, age, gender, symptoms, past_medical_history)

    # Display the report
    st.subheader("📢 AI Diagnosis Report")
    st.write(f"👤 **Patient Name:** {name}")
    st.write(f"🎂 **Age:** {age}")
    st.write(f"🚻 **Gender:** {gender}")
    st.write(f"🔍 **Symptoms:** {symptoms}")
    
    if past_medical_history:
        st.write(f"📜 **Past Medical History:** {past_medical_history}")

    st.write(f"🤖 **AI Diagnosis:** {response}")

    # Check for critical condition
    CRITICAL_KEYWORDS = ["stroke", "heart attack", "severe breathing difficulty", "unconscious", "chest pain"]
    
    if any(keyword in response.lower() for keyword in CRITICAL_KEYWORDS):
        st.error("🚨 Your symptoms may indicate a **critical condition**. Seek **IMMEDIATE medical attention!**")
        st.subheader("🏥 Emergency Contact Numbers (India)")
        for hospital, contact in INDIAN_HOSPITALS.items():
            st.write(f"📌 **{hospital}**: {contact}")

    # Generate medical report
    report_file = generate_report(name, age, gender, symptoms, past_medical_history, response)
    
    # Allow user to download report
    with open(report_file, "rb") as file:
        st.download_button(label="📄 Download Medical Report", data=file, file_name=report_file, mime="application/msword")

    # Store in chat history
    st.session_state.chat_history.append({
        "name": name,
        "age": age,
        "gender": gender,
        "symptoms": symptoms,
        "history": past_medical_history,
        "ai": response
    })

# Display previous diagnoses
st.subheader("📜 Previous Diagnoses")
for interaction in st.session_state.chat_history:
    st.write(f"👤 **{interaction['name']} ({interaction['age']} | {interaction['gender']})**")
    st.write(f"🔍 **Symptoms:** {interaction['symptoms']}")
    if interaction["history"]:
        st.write(f"📜 **Past History:** {interaction['history']}")
    st.write(f"🤖 **AI Diagnosis:** {interaction['ai']}")
    st.write("---")
